import streamlit as st
from docx import Document
from openai import OpenAI
import tempfile
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

def redaguj_opis(tekst, element="element budynku"):
    if not tekst or not tekst.strip():
        return ""

    prompt = f"""
Uredaguj poniższy opis z przeglądu budowlanego.

Zasady:
- język polski
- styl techniczny, rzeczowy i zwięzły
- popraw literówki, interpunkcję i szyk zdań
- usuń zbędne powtórzenia
- nie zmieniaj sensu wypowiedzi
- nie dopisuj informacji, których nie ma w tekście
- zachowaj informacje o uszkodzeniach, zużyciu, zawilgoceniu, pęknięciach, korozji itp.
- zwróć tylko gotowy opis, bez komentarza

Element: {element}

Tekst:
{tekst}
"""

    response = client.responses.create(
        model="gpt-4.1-mini",
        input=prompt
    )

    return response.output_text.strip()

st.set_page_config(page_title="Przegląd budynku", layout="wide")

st.title("Aplikacja do przeglądu budynku")
st.write("Formularz do wypełniania elementów przeglądu.")

WSZYSTKIE_ELEMENTY = [
    "Fundament",
    "Konstrukcja nośna",
    "Konstrukcja stropów",
    "Dach",
    "Elewacja",
    "Stolarka okienna",
    "Instalacje",
    "Inne",
]

if "elementy" not in st.session_state:
    st.session_state.elementy = []

if "opis_key" not in st.session_state:
    st.session_state.opis_key = 0


def generuj_word(lista_elementow):
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

    document = Document()
    document.add_heading("Protokół przeglądu budynku", level=1)

    for i, e in enumerate(lista_elementow, start=1):
        tabela = document.add_table(rows=3, cols=5)
        tabela.style = "Table Grid"
        tabela.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Szerokości kolumn
        tabela.columns[0].width = Cm(1.2)
        tabela.columns[1].width = Cm(4.5)
        tabela.columns[2].width = Cm(6.5)
        tabela.columns[3].width = Cm(3.0)
        tabela.columns[4].width = Cm(3.0)

        # ===== WIERSZ 1 =====
        row1 = tabela.rows[0].cells
        row1[0].text = "Lp."
        row1[1].text = "Nazwa elementu:"
        row1[2].text = str(e["element"])
        row1[2].merge(row1[4])

        # ===== WIERSZ 2 =====
        row2 = tabela.rows[1].cells
        row2[0].text = str(i)
        row2[1].text = "Opis:"
        row2[2].text = str(e["opis"])
        row2[0].merge(tabela.rows[2].cells[0])
        row2[2].merge(row2[4])

        # ===== WIERSZ 3 =====
        row3 = tabela.rows[2].cells
        row3[1].text = "Stan techniczny:"
        row3[2].text = str(e["stan"])
        row3[3].text = "Zużycie:"
        row3[4].text = f"{e['zuzycie']}%"

        # Wyrównania
        for row in tabela.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Opis do lewej
        for p in tabela.rows[1].cells[2].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        document.add_paragraph("")

    nazwa = "protokol_przegladu.docx"
    document.save(nazwa)
    return nazwa


# Elementy jeszcze nieopisane
wybrane_elementy = [e["element"] for e in st.session_state.elementy]
dostepne_elementy = [e for e in WSZYSTKIE_ELEMENTY if e not in wybrane_elementy]

if dostepne_elementy:
    col1, col2, col3 = st.columns([2, 1, 1])

    with col1:
        element = st.selectbox("Nazwa elementu", dostepne_elementy)

    with col2:
        stan = st.selectbox(
            "Stan techniczny",
            [
                "dobry",
                "zadowalający",
                "średni",
                "niżej średniego (lichy)",
                "zły",
            ],
        )

    if stan == "dobry":
        opcje_zuzycia = [0, 5, 10, 15]
    elif stan == "zadowalający":
        opcje_zuzycia = [16, 20, 25, 30]
    elif stan == "średni":
        opcje_zuzycia = [31, 35, 40, 45, 50]
    elif stan == "niżej średniego (lichy)":
        opcje_zuzycia = [51, 55, 60, 65, 70]
    else:
        opcje_zuzycia = [71, 75, 80, 85, 90, 95, 100]

    with col3:
        zuzycie = st.selectbox("Zużycie (%)", opcje_zuzycia)

st.subheader("Nagrywanie opisu")

auto_redakcja = st.checkbox("Automatycznie redaguj opis", value=True)

audio = st.audio_input("🎤 Nagraj opis głosowy")

if audio is not None:
    st.audio(audio)

    if st.button("Przepisz nagranie"):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp_file:
            tmp_file.write(audio.read())
            tmp_path = tmp_file.name

        with open(tmp_path, "rb") as f:
            transcript = client.audio.transcriptions.create(
                model="gpt-4o-mini-transcribe",
                file=f
            )

        tekst = transcript.text

        if auto_redakcja:
            tekst = redaguj_opis(tekst, st.session_state.opis_key)

        poprzedni_opis = st.session_state.get(
            f"opis_{st.session_state.opis_key}", ""
        ).strip()

        if poprzedni_opis:
            nowy_opis = poprzedni_opis + " " + tekst
        else:
            nowy_opis = tekst

        st.session_state[f"opis_{st.session_state.opis_key}"] = nowy_opis

        st.rerun()

    opis = st.text_area(
        "Opis elementu",
        height=200,
        key=f"opis_{st.session_state.opis_key}"
    )

    if st.button("➕ Dodaj element do tabeli", use_container_width=True):
        if opis.strip():
            st.session_state.elementy.append(
                {
                    "element": element,
                    "opis": opis,
                    "stan": stan,
                    "zuzycie": zuzycie,
                }
            )
            st.session_state.opis_key += 1
            st.success("Element zapisany")
            st.rerun()
        else:
            st.warning("Najpierw wpisz opis elementu")
else:
    st.info("Wszystkie elementy z listy zostały już opisane.")

st.subheader("Wprowadzone elementy")

if st.session_state.elementy:
    for i, e in enumerate(st.session_state.elementy, start=1):
        with st.expander(f"{i}. {e['element']}"):
            st.write("**Stan techniczny:**", e["stan"])
            st.write("**Zużycie:**", f"{e['zuzycie']} %")
            st.write("**Opis:**")
            st.write(e["opis"])
else:
    st.info("Brak wprowadzonych elementów")

col_a, col_b = st.columns(2)

with col_a:
    if st.button("Wyczyść wszystkie elementy", use_container_width=True):
        st.session_state.elementy = []
        st.session_state.opis_key += 1
        st.rerun()

with col_b:
    if st.session_state.elementy:
        plik = generuj_word(st.session_state.elementy)
        with open(plik, "rb") as f:
            st.download_button(
                label="Pobierz dokument Word",
                data=f,
                file_name=plik,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )